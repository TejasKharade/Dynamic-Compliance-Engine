import json
import csv
import pandas as pd
import pypdf
import pdfplumber
import docx

def _chunk_text(text: str, max_chars: int = 8000) -> list[str]:
    # A simple but robust chunking by paragraphs (\n\n) or single newlines
    chunks = []
    current_chunk = ""
    
    paragraphs = text.split("\n\n")
    for para in paragraphs:
        if len(current_chunk) + len(para) + 2 <= max_chars:
            current_chunk += para + "\n\n"
        else:
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            # If a single paragraph is longer than max_chars, we split by lines
            if len(para) > max_chars:
                lines = para.split("\n")
                current_chunk = ""
                for line in lines:
                    if len(current_chunk) + len(line) + 1 <= max_chars:
                        current_chunk += line + "\n"
                    else:
                        if current_chunk.strip():
                            chunks.append(current_chunk.strip())
                        current_chunk = line + "\n"
            else:
                current_chunk = para + "\n\n"
                
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
        
    return chunks

def extract_and_chunk_rules_file(filepath: str, max_chars: int = 8000) -> list[str]:
    ext = filepath.lower().split('.')[-1]
    
    text = ""
    if ext == "txt":
        with open(filepath, 'r', encoding='utf-8') as f:
            text = f.read()
            
    elif ext == "docx":
        doc = docx.Document(filepath)
        text = "\n\n".join([p.text for p in doc.paragraphs])
        # Add tables too
        for table in doc.tables:
            text += "\n\n"
            for row in table.rows:
                text += " | ".join([cell.text.replace("\n", " ").strip() for cell in row.cells]) + "\n"
                
    elif ext == "pdf":
        try:
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                        
                    # Extract tables
                    tables = page.extract_tables()
                    for table in tables:
                        text += "\n\n--- TABLE START ---\n"
                        for row in table:
                            # Filter out None cells
                            clean_row = [str(cell).replace("\n", " ").strip() if cell is not None else "" for cell in row]
                            text += " | ".join(clean_row) + "\n"
                        text += "--- TABLE END ---\n\n"
        except Exception as e:
            print(f"pdfplumber failed: {e}. Falling back to pypdf.")
            reader = pypdf.PdfReader(filepath)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
                    
    elif ext in ["csv", "xlsx"]:
        if ext == "csv":
            df = pd.read_csv(filepath)
        else:
            df = pd.read_excel(filepath)
        text = df.to_csv(index=False)
        
    else:
        raise ValueError(f"Unsupported file format: {ext}")

    return _chunk_text(text, max_chars)


def parse_inventory_file(filepath: str) -> list[dict]:
    ext = filepath.lower().split('.')[-1]
    
    if ext == "json":
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if not isinstance(data, list):
                data = [data] # Allow single JSON objects
            return data
            
    elif ext in ["csv", "xlsx"]:
        if ext == "csv":
            df = pd.read_csv(filepath)
        else:
            df = pd.read_excel(filepath)
            
        inventory = []
        for _, row in df.iterrows():
            raw_record = {}
            for col in df.columns:
                val = row[col]
                if pd.notna(val) and str(val).strip():
                    raw_record[col] = str(val).strip()
            if raw_record:
                inventory.append(raw_record)
        return inventory
        
    else:
        raise ValueError(f"Unsupported inventory format: {ext}")
